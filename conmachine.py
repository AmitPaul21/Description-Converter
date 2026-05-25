import streamlit as st
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_BREAK, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from io import BytesIO
import re
import pandas as pd

# ---------------- PAGE CONFIG ----------------
st.set_page_config(page_title="Description Converter", layout="centered")

# ---------------- Safe Whole Word Replacement ----------------
def replace_words_safe(text, replacement_dict):
    for key, val in replacement_dict.items():
        pattern = r'\b{}\b'.format(re.escape(str(key)))
        text = re.sub(pattern, str(val), text)
    return text

# ---------------- Clean Special Formatting ----------------
def clean_text(text):
    while '  ' in text:
        text = text.replace('  ', ' ')

    text = re.sub(r',\s*', ', ', text)
    text = re.sub(r'\(', ' (', text)

    while '  ' in text:
        text = text.replace('  ', ' ')

    text = re.sub(
        r'\b([A-Z]+)\s*\(',
        lambda m: m.group(1).capitalize() + ' (',
        text
    )

    return text

# ---------------- Remove Extra Empty Paragraphs ----------------
def remove_extra_empty_paragraphs(doc):
    paras = doc.paragraphs
    i = 0

    while i < len(paras) - 1:
        if paras[i].text.strip() == '' and paras[i + 1].text.strip() == '':
            p = paras[i]._element
            p.getparent().remove(p)
            paras = doc.paragraphs
            i -= 1
        i += 1

    return doc

# ---------------- Styled Line Break Helper ----------------
def add_styled_break(para, break_type=WD_BREAK.LINE):
    r = para.add_run("")
    r.font.name = "Times New Roman"
    r.font.size = Pt(10)

    rPr = r._element.get_or_add_rPr()

    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:cs"), "Times New Roman")
    rFonts.set(qn("w:eastAsia"), "Times New Roman")

    rPr.append(rFonts)

    r.add_break(break_type)

    return r

# ---------------- Highlight Helper ----------------
def highlight_run_safe(run, color="yellow"):
    if color == "yellow":
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    elif color == "red":
        run.font.highlight_color = WD_COLOR_INDEX.RED

# ---------------- Parse Depth Interval ----------------
def parse_depth_interval(text):
    match = re.match(r"^\s*(\d+)-(\d+)", text)

    if match:
        return int(match.group(1)), int(match.group(2)), match.group(0)

    return None

# ---------------- Check Percentages ----------------
def check_percentages_in_text(text):
    percentages = [int(x) for x in re.findall(r"\((\d+)%\)", text)]

    if percentages and sum(percentages) != 100:
        return True

    return False

# ---------------- Set paragraph text with highlighting ----------------
def set_para_text_with_highlight(para, text, prev_end=None):

    depth_result = parse_depth_interval(text)

    highlight_depth = False

    if depth_result:
        start, end, depth_text = depth_result

        if prev_end is not None and start != prev_end:
            highlight_depth = True

        prev_end = end

    else:
        depth_text = None

    highlight_percentage = check_percentages_in_text(text)

    # ---------------- If line starts with depth ----------------
    if depth_text and text.startswith(depth_text):

        r_depth = para.add_run(depth_text)
        r_depth.font.name = "Times New Roman"
        r_depth.font.size = Pt(10)

        if highlight_depth:
            highlight_run_safe(r_depth, "yellow")

        add_styled_break(para)

        rest_text = text[len(depth_text):].lstrip()

        if rest_text:
            r_rest = para.add_run(rest_text)
            r_rest.font.name = "Times New Roman"
            r_rest.font.size = Pt(10)

            if highlight_percentage:
                highlight_run_safe(r_rest, "red")

    # ---------------- No depth ----------------
    else:
        r = para.add_run(text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(10)

        if highlight_percentage:
            highlight_run_safe(r, "red")

    return prev_end

# ---------------- Load Excel Dictionary ----------------
@st.cache_data(ttl=3600)
def load_replacement_dict():

    df = pd.read_excel(
        "replacement_dict.xlsx",
        engine="openpyxl"
    )

    df = df.dropna(subset=['Find'])

    return dict(zip(df['Find'], df['Replace With']))

# ---------------- Document Formatting ----------------
def format_document(doc, replacement_dict):

    try:
        normal = doc.styles['Normal']

        normal.font.name = 'Times New Roman'
        normal.font.size = Pt(10)

        normal.paragraph_format.line_spacing = 1.5
        normal.paragraph_format.space_before = Pt(0)
        normal.paragraph_format.space_after = Pt(0)

    except Exception:
        normal = None

    prev_end = None

    # ---------------- Paragraphs ----------------
    for para in doc.paragraphs:

        raw = para.text

        text = replace_words_safe(raw, replacement_dict)

        text = clean_text(text).strip()

        if text:
            text = re.sub(r'[^A-Za-z0-9]+$', '', text)
            text += "."

        for run in list(para.runs):
            run.text = ''

        try:
            para._p.clear_content()
        except Exception:
            pass

        prev_end = set_para_text_with_highlight(
            para,
            text,
            prev_end
        )

        if normal is not None:
            para.style = normal

        para.paragraph_format.line_spacing = 1.5
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)

    doc = remove_extra_empty_paragraphs(doc)

    # ---------------- Tables ----------------
    for table in doc.tables:

        for row in table.rows:

            for cell in row.cells:

                for para in cell.paragraphs:

                    raw = para.text

                    text = replace_words_safe(raw, replacement_dict)

                    text = clean_text(text).strip()

                    for run in list(para.runs):
                        run.text = ''

                    try:
                        para._p.clear_content()
                    except Exception:
                        pass

                    prev_end = set_para_text_with_highlight(
                        para,
                        text,
                        prev_end
                    )

                    if normal is not None:
                        para.style = normal

                    para.paragraph_format.line_spacing = 1.5
                    para.paragraph_format.space_before = Pt(0)
                    para.paragraph_format.space_after = Pt(0)

    # ---------------- Headers and Footers ----------------
    for section in doc.sections:

        for hf in (section.header, section.footer):

            for para in hf.paragraphs:

                raw = para.text

                text = replace_words_safe(raw, replacement_dict)

                text = clean_text(text).strip()

                for run in list(para.runs):
                    run.text = ''

                try:
                    para._p.clear_content()
                except Exception:
                    pass

                prev_end = set_para_text_with_highlight(
                    para,
                    text,
                    prev_end
                )

                if normal is not None:
                    para.style = normal

                para.paragraph_format.line_spacing = 1.5
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(0)

    # ---------------- Page Margins ----------------
    for section in doc.sections:

        section.top_margin = Cm(2.29)
        section.bottom_margin = Cm(1.27)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    return doc

# ---------------- Custom Header Formatter ----------------
def format_header(doc, well_name):

    for section in doc.sections:

        header = section.header

        for para in header.paragraphs:
            p = para._element
            p.getparent().remove(p)

        # ---------------- Well Name ----------------
        para1 = header.add_paragraph()
        para1.alignment = 1
        para1.paragraph_format.line_spacing = 1.0

        run1 = para1.add_run(well_name.upper())

        run1.font.name = "Times New Roman"
        run1.font.size = Pt(10)

        # ---------------- Space ----------------
        header.add_paragraph().paragraph_format.line_spacing = 1.0

        # ---------------- Title ----------------
        para2 = header.add_paragraph()
        para2.alignment = 1
        para2.paragraph_format.line_spacing = 1.0

        run2 = para2.add_run("SAMPLE DESCRIPTIONS")

        run2.font.name = "Times New Roman"
        run2.font.size = Pt(10)
        run2.underline = True

        # ---------------- Space ----------------
        header.add_paragraph().paragraph_format.line_spacing = 1.0

        # ---------------- Depth ----------------
        para3 = header.add_paragraph()
        para3.alignment = 0
        para3.paragraph_format.line_spacing = 1.0

        run3 = para3.add_run("Depth (m)")

        run3.font.name = "Times New Roman"
        run3.font.size = Pt(10)

        header.add_paragraph().paragraph_format.line_spacing = 1.0

    # ---------------- Remove Footer ----------------
    for section in doc.sections:

        footer = section.footer

        for para in footer.paragraphs:
            p = para._element
            p.getparent().remove(p)

    return doc

# ---------------- STREAMLIT APP ----------------
st.title("Description Converter 📝")

st.markdown("""
### Instructions 📌
1. Create a Word document and insert the text file for the Build Section.
2. Go to Layout > Breaks > Page Break.
3. Insert the text file for the Horizontal Section.
4. Save the Word file.
5. Upload the saved file here.
6. Enter the well name.
7. Download the processed file.
8. Yellow highlight = depth issue.
9. Red highlight = percentage issue.
""")

uploaded_file = st.file_uploader(
    "Upload your Word file (.docx)",
    type=["docx"]
)

well_name = st.text_input(
    "Enter Well Name (will appear in header)",
    ""
)

# ---------------- PROCESS BUTTON ----------------
if uploaded_file is not None and st.button("Process File"):

    try:

        # ---------------- Load Excel ONLY when needed ----------------
        replacement_dict = load_replacement_dict()

        # ---------------- Load Document ----------------
        doc = Document(uploaded_file)

        # ---------------- Format Document ----------------
        doc = format_document(doc, replacement_dict)

        # ---------------- Add Header ----------------
        if well_name.strip():
            doc = format_header(doc, well_name)

        # ---------------- Save Output ----------------
        buffer = BytesIO()

        doc.save(buffer)

        buffer.seek(0)

        st.success("✅ File processed successfully!")

        st.download_button(
            label="Download Final File",
            data=buffer,
            file_name="formatted_checked_output.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    except Exception as e:

        st.error(f"❌ Error: {e}")
